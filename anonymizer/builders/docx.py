"""Reconstructor de documentos Word (.docx) anonimizados."""

from io import BytesIO
from typing import List

from docx import Document

from anonymizer.builders.base import BaseBuilder
from anonymizer.extractors.base import ExtractedText


class DocxBuilder(BaseBuilder):
    """Reconstruye un documento Word reemplazando el texto por su versión anonimizada."""

    def build(
        self,
        original_stream: BytesIO,
        extracted_texts: List[ExtractedText],
        redacted_texts: List[str],
    ) -> BytesIO:
        original_stream.seek(0)
        document = Document(original_stream)

        # Índice para recorrer los textos anonimizados
        idx = 0

        # Párrafos
        for para_idx, paragraph in enumerate(document.paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                if idx < len(extracted_texts):
                    loc = extracted_texts[idx].location
                    if loc[0] == "paragraph" and loc[1] == para_idx and loc[2] == run_idx:
                        run.text = redacted_texts[idx]
                        idx += 1

        # Tablas
        for table_idx, table in enumerate(document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if idx < len(extracted_texts):
                                loc = extracted_texts[idx].location
                                if (
                                    loc[0] == "table"
                                    and loc[1] == table_idx
                                    and loc[2] == row_idx
                                    and loc[3] == cell_idx
                                    and loc[4] == para_idx
                                    and loc[5] == run_idx
                                ):
                                    run.text = redacted_texts[idx]
                                    idx += 1

        # Headers y footers
        for section in document.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header is not None:
                    for para_idx, paragraph in enumerate(header.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if idx < len(extracted_texts):
                                loc = extracted_texts[idx].location
                                if loc[0] == "header" and loc[1] == para_idx and loc[2] == run_idx:
                                    run.text = redacted_texts[idx]
                                    idx += 1
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer is not None:
                    for para_idx, paragraph in enumerate(footer.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if idx < len(extracted_texts):
                                loc = extracted_texts[idx].location
                                if loc[0] == "footer" and loc[1] == para_idx and loc[2] == run_idx:
                                    run.text = redacted_texts[idx]
                                    idx += 1

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return output

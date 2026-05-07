"""Extractor de texto para documentos Word (.docx)."""

from io import BytesIO
from typing import List

from docx import Document

from anonymizer.extractors.base import BaseExtractor, ExtractedText


class DocxExtractor(BaseExtractor):
    """Extrae texto de archivos Word preservando la estructura de runs."""

    def extract(self, file_stream: BytesIO) -> List[ExtractedText]:
        document = Document(file_stream)
        extracted: List[ExtractedText] = []

        # Párrafos
        for para_idx, paragraph in enumerate(document.paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                if run.text:
                    extracted.append(
                        ExtractedText(
                            text=run.text,
                            location=("paragraph", para_idx, run_idx),
                        )
                    )

        # Tablas
        for table_idx, table in enumerate(document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if run.text:
                                extracted.append(
                                    ExtractedText(
                                        text=run.text,
                                        location=(
                                            "table",
                                            table_idx,
                                            row_idx,
                                            cell_idx,
                                            para_idx,
                                            run_idx,
                                        ),
                                    )
                                )

        # Headers y footers
        for section in document.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header is not None:
                    for para_idx, paragraph in enumerate(header.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if run.text:
                                extracted.append(
                                    ExtractedText(
                                        text=run.text,
                                        location=("header", para_idx, run_idx),
                                    )
                                )
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer is not None:
                    for para_idx, paragraph in enumerate(footer.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if run.text:
                                extracted.append(
                                    ExtractedText(
                                        text=run.text,
                                        location=("footer", para_idx, run_idx),
                                    )
                                )

        return extracted
